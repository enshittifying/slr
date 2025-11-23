#!/usr/bin/env python3
"""
Substantive R1 Cite Checker with GPT-4o Support Validation

Per R1 Member Handbook:
- Extract text propositions (sentences with footnote markers)
- Check if each citation substantively supports its proposition
- Generate edits/comments ([AA:] or [SE:]) for Word document
- Output structured JSON with validation results and suggested edits

Key Principle: "Assume the author is wrong until convinced otherwise"
"""

import json
import re
import os
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass, asdict
from docx import Document
from lxml import etree
from openai import OpenAI


@dataclass
class TextProposition:
    """Text claim with footnote reference."""
    paragraph_index: int
    sentence: str
    footnote_refs: List[int]  # Footnote numbers referenced


@dataclass
class SubstantiveCheck:
    """Result of substantive support check."""
    support_level: str  # "yes", "maybe", "no"
    confidence: float  # 0.0-1.0
    reasoning: str
    supported_elements: List[str]
    unsupported_elements: List[str]
    suggested_action: str  # "accept", "flag_aa", "flag_se", "reject"
    missing_context: str


@dataclass
class CitationWithSupport:
    """Citation with substantive support analysis."""
    citation_number: int
    citation_text: str
    text_proposition: str
    support_check: Optional[SubstantiveCheck] = None
    formatting_errors: List[Dict] = None
    edit: Optional[Dict] = None  # Suggested edit for Word doc

    def __post_init__(self):
        if self.formatting_errors is None:
            self.formatting_errors = []


@dataclass
class FootnoteWithProposition:
    """Footnote linked to its text proposition."""
    footnote_id: str
    footnote_number: int
    footnote_text: str
    text_proposition: str  # The sentence this footnote supports
    paragraph_index: int
    citations: List[CitationWithSupport] = None

    def __post_init__(self):
        if self.citations is None:
            self.citations = []


class SubstantiveR1Checker:
    """R1 substantive cite checker with GPT-4o validation."""

    def __init__(self, api_key: str = None):
        """Initialize with OpenAI API key."""
        # Get API key from environment or parameter
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            # Try to read from file
            api_key_file = Path.home() / ".openai_api_key"
            if api_key_file.exists():
                self.api_key = api_key_file.read_text().strip()

        if not self.api_key:
            raise ValueError("No OpenAI API key found. Set OPENAI_API_KEY env var or pass api_key parameter")

        self.client = OpenAI(api_key=self.api_key)
        self.model = "gpt-4o"  # Use GPT-4o for substantive checks
        self.total_tokens = 0
        self.total_cost = 0.0

    def extract_text_propositions(self, docx_path: Path) -> List[TextProposition]:
        """Extract text propositions (sentences with footnote markers)."""
        doc = Document(docx_path)
        propositions = []

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Find all footnote references in this paragraph
            # Pattern: superscript numbers or [number]
            footnote_pattern = r'(?:(?<=\s)|(?<=\.)|(?<=,))(\d+)(?=\s|$|,|\.)'

            # Split into sentences
            sentences = re.split(r'(?<=[.!?])\s+', text)

            for sentence in sentences:
                # Find footnote references in this sentence
                refs = []
                for match in re.finditer(footnote_pattern, sentence):
                    fn_num = int(match.group(1))
                    refs.append(fn_num)

                if refs:
                    propositions.append(TextProposition(
                        paragraph_index=para_idx,
                        sentence=sentence,
                        footnote_refs=refs
                    ))

        return propositions

    def extract_footnotes_with_ids(self, docx_path: Path) -> List[Dict]:
        """Extract footnotes with XML IDs."""
        doc = Document(docx_path)
        footnotes = []

        # Access footnotes through XML
        doc_part = doc.part
        footnotes_part = None

        for rel in doc_part.rels.values():
            if "footnotes" in rel.target_ref:
                footnotes_part = rel.target_part
                break

        if not footnotes_part:
            return footnotes

        # Parse XML
        footnotes_xml = footnotes_part.blob
        root = etree.fromstring(footnotes_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        footnote_counter = 0

        for footnote in root.findall('.//w:footnote', ns):
            fn_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')

            if fn_id is None or fn_id in ['-1', '0']:
                continue

            # Extract text
            footnote_text = ""
            for para in footnote.findall('.//w:p', ns):
                para_text = []
                for run in para.findall('.//w:r', ns):
                    for text_elem in run.findall('.//w:t', ns):
                        if text_elem.text:
                            para_text.append(text_elem.text)
                footnote_text += ''.join(para_text) + " "

            if footnote_text.strip():
                footnote_counter += 1
                footnotes.append({
                    'id': fn_id,
                    'number': footnote_counter,
                    'text': footnote_text.strip()
                })

        return footnotes

    def link_footnotes_to_propositions(self,
                                      propositions: List[TextProposition],
                                      footnotes: List[Dict]) -> List[FootnoteWithProposition]:
        """Link each footnote to its text proposition."""
        footnote_map = {fn['number']: fn for fn in footnotes}
        linked = []

        for prop in propositions:
            for fn_num in prop.footnote_refs:
                if fn_num in footnote_map:
                    fn = footnote_map[fn_num]
                    linked.append(FootnoteWithProposition(
                        footnote_id=fn['id'],
                        footnote_number=fn['number'],
                        footnote_text=fn['text'],
                        text_proposition=prop.sentence,
                        paragraph_index=prop.paragraph_index
                    ))

        return linked

    def parse_citations_in_footnote(self, footnote_text: str, proposition: str) -> List[CitationWithSupport]:
        """Parse individual citations within a footnote."""
        citations = []

        # Split on semicolons or periods followed by capital letter
        citation_parts = re.split(r';\s+(?=[A-Z])|\.(?=\s+[A-Z][a-z]+\s+v\.)', footnote_text)

        for i, citation_text in enumerate(citation_parts, 1):
            if citation_text.strip():
                citations.append(CitationWithSupport(
                    citation_number=i,
                    citation_text=citation_text.strip(),
                    text_proposition=proposition
                ))

        return citations

    def check_substantive_support(self,
                                  proposition: str,
                                  citation: str) -> SubstantiveCheck:
        """
        Use GPT-4o to check if citation supports proposition.

        Key Principle: Assume author is wrong until convinced otherwise.
        """
        system_prompt = """You are an expert legal research editor performing R1 cite checking for Stanford Law Review.

Your task is to rigorously evaluate whether a citation substantively supports a text proposition.

CRITICAL PRINCIPLE: Assume the author is wrong until you are convinced otherwise.

Be STRICT:
- Require DIRECT support, not just tangential relevance
- Flag if citation doesn't fully support ALL elements of the proposition
- Note missing context or qualifications
- Identify overstated claims

Respond with JSON only."""

        user_prompt = f"""Evaluate if this citation supports the proposition.

TEXT PROPOSITION (claim in the article):
{proposition}

CITATION:
{citation}

Respond with JSON:
{{
  "support_level": "yes" | "maybe" | "no",
  "confidence": 0.0-1.0,
  "reasoning": "detailed explanation",
  "supported_elements": ["list of claims that ARE supported"],
  "unsupported_elements": ["list of claims that are NOT supported"],
  "suggested_action": "accept" | "flag_aa" | "flag_se" | "reject",
  "missing_context": "what's missing or unclear"
}}

Where:
- "accept": Citation fully supports proposition
- "flag_aa": Flag for author attention [AA:] - minor issue
- "flag_se": Flag for senior editor [SE:] - significant issue
- "reject": Citation does not support proposition

Be rigorous. If in doubt, flag it."""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                response_format={"type": "json_object"},
                temperature=0.1  # Low temperature for consistency
            )

            # Track usage
            self.total_tokens += response.usage.total_tokens
            self.total_cost += (response.usage.prompt_tokens * 0.0025 / 1000 +
                              response.usage.completion_tokens * 0.01 / 1000)

            result = json.loads(response.choices[0].message.content)

            return SubstantiveCheck(
                support_level=result.get('support_level', 'unknown'),
                confidence=float(result.get('confidence', 0.0)),
                reasoning=result.get('reasoning', ''),
                supported_elements=result.get('supported_elements', []),
                unsupported_elements=result.get('unsupported_elements', []),
                suggested_action=result.get('suggested_action', 'flag_se'),
                missing_context=result.get('missing_context', '')
            )

        except Exception as e:
            print(f"  ⚠️  GPT-4o call failed: {e}")
            return SubstantiveCheck(
                support_level="error",
                confidence=0.0,
                reasoning=f"API error: {str(e)}",
                supported_elements=[],
                unsupported_elements=[],
                suggested_action="flag_se",
                missing_context="Could not validate"
            )

    def generate_edit_suggestion(self, citation: CitationWithSupport) -> Dict:
        """Generate edit/comment for Word document."""
        if not citation.support_check:
            return None

        check = citation.support_check

        if check.suggested_action == "accept":
            return None  # No edit needed

        elif check.suggested_action == "flag_aa":
            comment_text = f"[AA: {check.reasoning}]"
            if check.missing_context:
                comment_text += f" Missing context: {check.missing_context}"

            return {
                "type": "comment",
                "marker": "AA",
                "text": comment_text,
                "severity": "minor",
                "location": "footnote",
                "footnote_number": None  # Will be filled by caller
            }

        elif check.suggested_action == "flag_se":
            comment_text = f"[SE: {check.reasoning}]"
            if check.unsupported_elements:
                comment_text += f" Unsupported: {', '.join(check.unsupported_elements)}"

            return {
                "type": "comment",
                "marker": "SE",
                "text": comment_text,
                "severity": "major",
                "location": "footnote",
                "footnote_number": None
            }

        elif check.suggested_action == "reject":
            return {
                "type": "edit",
                "action": "delete_or_revise",
                "text": f"[SE: Citation does not support proposition. {check.reasoning}]",
                "severity": "critical",
                "location": "footnote",
                "footnote_number": None
            }

        return None

    def check_article_substantive(self, docx_path: Path, sample_size: int = None) -> Dict:
        """
        Perform substantive R1 cite checking on entire article.

        Args:
            docx_path: Path to Word document
            sample_size: If set, only check first N footnotes (for testing)

        Returns:
            Dict with structured results and edits
        """
        print("="*80)
        print("SUBSTANTIVE R1 CITE CHECKING (GPT-4o Validation)")
        print("="*80)

        # Extract text propositions
        print(f"\n1. Extracting text propositions...")
        propositions = self.extract_text_propositions(docx_path)
        print(f"   ✅ Found {len(propositions)} text propositions with footnotes")

        # Extract footnotes
        print(f"\n2. Extracting footnotes...")
        footnotes = self.extract_footnotes_with_ids(docx_path)
        print(f"   ✅ Extracted {len(footnotes)} footnotes")

        # Link footnotes to propositions
        print(f"\n3. Linking footnotes to text propositions...")
        linked = self.link_footnotes_to_propositions(propositions, footnotes)
        print(f"   ✅ Linked {len(linked)} footnote-proposition pairs")

        # Apply sample size limit if set
        if sample_size:
            linked = linked[:sample_size]
            print(f"   ℹ️  Sample mode: Checking first {sample_size} footnotes")

        # Parse citations and check support
        print(f"\n4. Checking substantive support with GPT-4o...")
        print(f"   Model: {self.model}")
        print(f"   Footnotes to check: {len(linked)}")

        total_citations = 0
        checks_performed = 0

        for i, fn in enumerate(linked, 1):
            print(f"\n   [{i}/{len(linked)}] Footnote {fn.footnote_number}...")

            # Parse citations
            fn.citations = self.parse_citations_in_footnote(
                fn.footnote_text,
                fn.text_proposition
            )
            total_citations += len(fn.citations)

            # Check each citation
            for j, citation in enumerate(fn.citations, 1):
                print(f"      Citation {j}/{len(fn.citations)}: Checking support...")

                # Perform substantive check
                citation.support_check = self.check_substantive_support(
                    citation.text_proposition,
                    citation.citation_text
                )
                checks_performed += 1

                # Generate edit suggestion
                citation.edit = self.generate_edit_suggestion(citation)
                if citation.edit:
                    citation.edit['footnote_number'] = fn.footnote_number

                # Print result
                check = citation.support_check
                support_icon = "✅" if check.support_level == "yes" else "⚠️" if check.support_level == "maybe" else "❌"
                print(f"      {support_icon} Support: {check.support_level} ({check.confidence:.0%}) - {check.suggested_action}")

        print(f"\n   ✅ Completed {checks_performed} substantive checks")

        # Compile results
        results = {
            'total_propositions': len(propositions),
            'total_footnotes_checked': len(linked),
            'total_citations_checked': total_citations,
            'total_support_checks': checks_performed,
            'api_usage': {
                'model': self.model,
                'total_tokens': self.total_tokens,
                'estimated_cost': f"${self.total_cost:.2f}"
            },
            'footnotes': [asdict(fn) for fn in linked],
            'summary': self._generate_summary(linked)
        }

        return results

    def _generate_summary(self, footnotes: List[FootnoteWithProposition]) -> Dict:
        """Generate summary statistics."""
        summary = {
            'support_levels': {'yes': 0, 'maybe': 0, 'no': 0, 'error': 0},
            'suggested_actions': {'accept': 0, 'flag_aa': 0, 'flag_se': 0, 'reject': 0},
            'edits_needed': 0,
            'citations_by_support': {}
        }

        for fn in footnotes:
            for citation in fn.citations:
                if citation.support_check:
                    check = citation.support_check
                    summary['support_levels'][check.support_level] = \
                        summary['support_levels'].get(check.support_level, 0) + 1
                    summary['suggested_actions'][check.suggested_action] = \
                        summary['suggested_actions'].get(check.suggested_action, 0) + 1

                if citation.edit:
                    summary['edits_needed'] += 1

        return summary

    def print_report(self, results: Dict):
        """Print formatted substantive checking report."""
        print("\n" + "="*80)
        print("SUBSTANTIVE R1 CHECKING REPORT")
        print("="*80)

        print(f"\n📊 COVERAGE:")
        print(f"   • {results['total_propositions']} text propositions")
        print(f"   • {results['total_footnotes_checked']} footnotes checked")
        print(f"   • {results['total_citations_checked']} citations analyzed")

        print(f"\n🤖 GPT-4o VALIDATION:")
        summary = results['summary']
        print(f"   • {summary['support_levels']['yes']} citations FULLY support")
        print(f"   • {summary['support_levels']['maybe']} citations PARTIALLY support")
        print(f"   • {summary['support_levels']['no']} citations DO NOT support")
        if summary['support_levels']['error'] > 0:
            print(f"   • {summary['support_levels']['error']} validation errors")

        print(f"\n📝 SUGGESTED ACTIONS:")
        print(f"   • {summary['suggested_actions']['accept']} citations accepted")
        print(f"   • {summary['suggested_actions']['flag_aa']} [AA:] author flags")
        print(f"   • {summary['suggested_actions']['flag_se']} [SE:] senior editor flags")
        print(f"   • {summary['suggested_actions']['reject']} citations to reject")

        print(f"\n✏️  EDITS NEEDED:")
        print(f"   • {summary['edits_needed']} edits/comments to add to document")

        print(f"\n💰 API USAGE:")
        print(f"   • Model: {results['api_usage']['model']}")
        print(f"   • Tokens: {results['api_usage']['total_tokens']:,}")
        print(f"   • Cost: {results['api_usage']['estimated_cost']}")

        # Show sample problematic citations
        print(f"\n🔍 SAMPLE ISSUES (First 10):")
        issue_count = 0
        for fn in results['footnotes']:
            for citation in fn['citations']:
                if citation.get('edit'):
                    issue_count += 1
                    if issue_count <= 10:
                        check = citation.get('support_check', {})
                        print(f"\n{issue_count}. Footnote {fn['footnote_number']}, Citation {citation['citation_number']}")
                        print(f"   Proposition: {citation['text_proposition'][:80]}...")
                        print(f"   Support: {check.get('support_level', 'unknown')} ({check.get('confidence', 0):.0%})")
                        print(f"   Issue: {check.get('reasoning', 'Unknown')[:100]}...")
                        print(f"   Action: {check.get('suggested_action', 'unknown')}")

        print("\n" + "="*80)


def main():
    """Main function."""
    # Check for API key
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        api_key_file = Path.home() / ".openai_api_key"
        if api_key_file.exists():
            api_key = api_key_file.read_text().strip()

    if not api_key:
        print("❌ No OpenAI API key found!")
        print("   Set OPENAI_API_KEY environment variable or create ~/.openai_api_key file")
        return 1

    # Initialize checker
    checker = SubstantiveR1Checker(api_key=api_key)

    # Check Sanders article
    article_path = Path("/home/user/slr/sp_machine/Sanders_PreSP_PostEEFormatting.docx")

    if not article_path.exists():
        print(f"❌ Article not found at {article_path}")
        return 1

    # Run substantive check (sample first 5 footnotes for testing)
    results = checker.check_article_substantive(article_path, sample_size=5)

    # Print report
    checker.print_report(results)

    # Save results to JSON
    output_path = Path(__file__).parent.parent / "output" / "sanders_substantive_check.json"
    output_path.parent.mkdir(exist_ok=True)

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)

    print(f"\n💾 Full results saved to: {output_path}")
    print("="*80)

    return 0


if __name__ == '__main__':
    exit(main())
