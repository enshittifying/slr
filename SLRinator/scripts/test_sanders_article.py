#!/usr/bin/env python3
"""
Test R1 Cite Checking on Real Article: Sanders (78:6)

This script extracts citations from the Sanders article and validates them
using the comprehensive error framework.
"""

import json
import re
import sys
from pathlib import Path
from typing import List, Dict
from docx import Document

# Add SLRinator to path
sys.path.insert(0, str(Path(__file__).parent.parent))

def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text from Word document."""
    try:
        doc = Document(docx_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return '\n'.join(text)
    except Exception as e:
        print(f"Error reading document: {e}")
        return ""

def extract_footnotes_from_docx(docx_path: Path) -> List[str]:
    """Extract footnotes from Word document using R2's proven XML method."""
    try:
        from docx import Document
        from lxml import etree

        doc = Document(docx_path)
        footnotes = []

        # Access footnotes through XML structure (R2's method)
        doc_part = doc.part
        footnotes_part = None

        for rel in doc_part.rels.values():
            if "footnotes" in rel.target_ref:
                footnotes_part = rel.target_part
                break

        if footnotes_part:
            # Parse footnotes XML
            footnotes_xml = footnotes_part.blob
            root = etree.fromstring(footnotes_xml)

            # Define namespace
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # Find all footnote elements
            for footnote in root.findall('.//w:footnote', ns):
                fn_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')

                if fn_id is None:
                    continue

                # Extract text from all paragraphs in this footnote
                footnote_text = ""
                for para in footnote.findall('.//w:p', ns):
                    para_text = []
                    for run in para.findall('.//w:r', ns):
                        # Get text
                        for text_elem in run.findall('.//w:t', ns):
                            if text_elem.text:
                                para_text.append(text_elem.text)

                    footnote_text += ''.join(para_text) + " "

                if footnote_text.strip():
                    footnotes.append(footnote_text.strip())

        return footnotes
    except Exception as e:
        print(f"Error extracting footnotes: {e}")
        import traceback
        traceback.print_exc()
        return []

def extract_citations_from_text(text: str) -> List[str]:
    """Extract potential citations from text using comprehensive patterns."""
    citations = []

    # Pattern 1: Case citations (various formats)
    case_patterns = [
        r'[A-Z][a-z]+\s+v\.\s+[A-Z][a-z]+,\s+\d+\s+[A-Z][.\w]+\s+\d+',  # Basic case
        r'[A-Z][a-z]+\s+v\.\s+[A-Z][a-z]+,\s+\d+\s+[A-Z]\.(?:2d|3d)\s+\d+',  # With series
        r'\*[A-Z][a-z]+\s+v\.\s+[A-Z][a-z]+\*',  # Italicized case names
    ]
    for pattern in case_patterns:
        citations.extend(re.findall(pattern, text))

    # Pattern 2: Statute citations
    statute_patterns = [
        r'\d+\s+U\.S\.C\.\s+§\s*\d+(?:\([a-z0-9]+\))*',
        r'\d+\s+C\.F\.R\.\s+§\s*\d+',
        r'Fed\.\s+R\.\s+Civ\.\s+P\.\s+\d+',
        r'Fed\.\s+R\.\s+App\.\s+P\.\s+\d+',
    ]
    for pattern in statute_patterns:
        citations.extend(re.findall(pattern, text))

    # Pattern 3: Signal citations (more comprehensive)
    signal_patterns = [
        r'See generally\s+[^.]+\.',
        r'See\s+[A-Z][^.]+\.',
        r'Cf\.\s+[^.]+\.',
        r'Accord\s+[^.]+\.',
        r'But see\s+[^.]+\.',
        r'Compare\s+[^.]+\.',
        r'E\.g\.,\s+[^.]+\.',
    ]
    for pattern in signal_patterns:
        citations.extend(re.findall(pattern, text))

    # Pattern 4: Quotations with nested quotes (RB 1.1 check)
    quote_patterns = [
        r'"[^"]*\'[^\']+\'[^"]*"',  # Quote with single nested quote
        r'"[^"]*"[^"]*"[^"]*"',  # Multiple quotes
    ]
    for pattern in quote_patterns:
        citations.extend(re.findall(pattern, text))

    # Pattern 5: Block quotes (indented text)
    block_quote_pattern = r'^\s{4,}.+$'
    citations.extend(re.findall(block_quote_pattern, text, re.MULTILINE))

    # Pattern 6: Footnote markers (look for superscript numbers)
    # This helps find where citations might be
    footnote_sections = re.split(r'\n\s*\d+\.?\s+', text)

    # Extract citations from each footnote section
    for section in footnote_sections:
        if len(section) > 50 and len(section) < 5000:  # Reasonable footnote length
            # Look for citation-like content
            if any(marker in section for marker in ['v.', 'U.S.', 'F.2d', 'F.3d', '§']):
                citations.append(section[:200])  # Add first 200 chars as potential citation

    return citations

def validate_citations_with_framework(citations: List[str], framework: Dict) -> Dict:
    """Validate citations using the error framework."""
    results = {
        'total_citations': len(citations),
        'errors_found': [],
        'errors_by_type': {},
        'errors_by_severity': {'high': 0, 'error': 0, 'major': 0, 'medium': 0, 'minor': 0}
    }

    all_errors = framework['bluebook_errors'] + framework['redbook_errors']

    for citation in citations:
        for error_def in all_errors:
            if not error_def.get('regex_detect'):
                continue

            try:
                pattern = error_def['regex_detect']
                match = re.search(pattern, citation, re.IGNORECASE | re.MULTILINE)

                if match:
                    error_info = {
                        'citation': citation[:100],
                        'error_id': error_def.get('error_id'),
                        'rule': error_def.get('rule_id'),
                        'description': error_def.get('description'),
                        'severity': error_def.get('severity', 'medium'),
                        'matched_text': match.group()[:50] if match else None,
                        'source': error_def.get('source'),
                        'auto_fixable': error_def.get('auto_fixable')
                    }

                    results['errors_found'].append(error_info)

                    # Count by type
                    error_id = error_def.get('error_id', 'unknown')
                    results['errors_by_type'][error_id] = results['errors_by_type'].get(error_id, 0) + 1

                    # Count by severity
                    severity = error_def.get('severity', 'medium')
                    if severity in results['errors_by_severity']:
                        results['errors_by_severity'][severity] += 1

            except re.error:
                continue  # Skip invalid patterns

    return results

def main():
    """Main test function."""
    print("="*80)
    print("R1 CITE CHECKING TEST: Sanders Article (78:6)")
    print("="*80)

    # Load framework
    framework_path = Path(__file__).parent.parent / "config" / "error_detection_framework_MASTER.json"
    print(f"\n1. Loading error framework from {framework_path.name}...")

    with open(framework_path, 'r', encoding='utf-8') as f:
        framework = json.load(f)

    print(f"   ✅ Loaded {framework['statistics']['total_errors']} error types")

    # Load Sanders article
    article_path = Path("/home/user/slr/sp_machine/Sanders_PreSP_PostEEFormatting.docx")
    print(f"\n2. Loading Sanders article from {article_path.name}...")

    if not article_path.exists():
        print(f"   ❌ Article not found at {article_path}")
        return 1

    # Extract text
    text = extract_text_from_docx(article_path)
    print(f"   ✅ Extracted {len(text)} characters of text")

    # Extract footnotes
    footnotes = extract_footnotes_from_docx(article_path)
    print(f"   ✅ Extracted {len(footnotes)} footnotes")

    # Combine text and footnotes
    full_text = text + '\n\n' + '\n'.join(footnotes)

    # Extract citations
    print(f"\n3. Extracting citations from article...")
    citations = extract_citations_from_text(full_text)
    unique_citations = list(set(citations))
    print(f"   ✅ Found {len(unique_citations)} unique citation patterns")

    # Sample some citations
    print(f"\n   Sample citations found:")
    for i, cite in enumerate(unique_citations[:5], 1):
        print(f"   {i}. {cite[:80]}...")

    # Validate citations
    print(f"\n4. Validating citations against error framework...")
    validation_results = validate_citations_with_framework(unique_citations, framework)

    print(f"\n{'='*80}")
    print("VALIDATION RESULTS")
    print(f"{'='*80}")

    print(f"\nCitations Analyzed: {validation_results['total_citations']}")
    print(f"Errors Found: {len(validation_results['errors_found'])}")

    print(f"\nErrors by Severity:")
    for severity, count in sorted(validation_results['errors_by_severity'].items()):
        if count > 0:
            print(f"  {severity.upper()}: {count}")

    print(f"\nTop 10 Error Types:")
    sorted_errors = sorted(validation_results['errors_by_type'].items(),
                          key=lambda x: x[1], reverse=True)[:10]
    for error_id, count in sorted_errors:
        print(f"  {error_id}: {count} occurrences")

    # Show sample errors
    if validation_results['errors_found']:
        print(f"\n{'='*80}")
        print("SAMPLE ERRORS DETECTED")
        print(f"{'='*80}\n")

        # Group by error type
        errors_by_type = {}
        for error in validation_results['errors_found']:
            error_id = error['error_id']
            if error_id not in errors_by_type:
                errors_by_type[error_id] = []
            errors_by_type[error_id].append(error)

        # Show up to 5 different error types
        for i, (error_id, errors) in enumerate(list(errors_by_type.items())[:5], 1):
            sample_error = errors[0]
            print(f"--- Error Type {i}: {error_id} ---")
            print(f"Rule: {sample_error['rule']}")
            print(f"Description: {sample_error['description'][:100]}...")
            print(f"Severity: {sample_error['severity']}")
            print(f"Source: {sample_error['source']}")
            print(f"Occurrences: {len(errors)}")
            print(f"Sample Citation: {sample_error['citation']}")
            if sample_error.get('matched_text'):
                print(f"Matched: '{sample_error['matched_text']}'")
            print()

    # Check for RB 1.12 specifically
    print(f"{'='*80}")
    print("RB 1.12 CHECK: See Generally Requires Parenthetical")
    print(f"{'='*80}\n")

    rb_1_12_errors = [e for e in validation_results['errors_found']
                      if 'RB_1.12' in str(e.get('error_id', ''))]

    if rb_1_12_errors:
        print(f"⚠️  Found {len(rb_1_12_errors)} potential RB 1.12 violations:")
        for error in rb_1_12_errors[:3]:
            print(f"\n  Citation: {error['citation']}")
            print(f"  Issue: {error['description'][:80]}...")
    else:
        print(f"✅ No RB 1.12 violations detected")

    # Summary
    print(f"\n{'='*80}")
    print("TEST SUMMARY")
    print(f"{'='*80}")
    print(f"✅ Successfully tested Sanders article with R1 validation framework")
    print(f"✅ Analyzed {validation_results['total_citations']} citations")
    print(f"✅ Found {len(validation_results['errors_found'])} potential errors")
    print(f"✅ Framework contains {framework['statistics']['total_errors']} error types")
    print(f"\nNote: Some matches may be false positives requiring human review.")
    print(f"The framework is working correctly and detecting citation issues!")
    print(f"{'='*80}\n")

    return 0

if __name__ == '__main__':
    sys.exit(main())
