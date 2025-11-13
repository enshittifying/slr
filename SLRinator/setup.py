#!/usr/bin/env python3
"""
Setup script for Stanford Law Review Editorial Automation System
"""

import os
import sys
import subprocess
import platform
from pathlib import Path
import shutil


def print_banner():
    """Print welcome banner"""
    banner = """
    ╔═══════════════════════════════════════════════════════════════╗
    ║                                                               ║
    ║     Stanford Law Review Editorial Automation System          ║
    ║                      (SLRinator v1.0)                        ║
    ║                                                               ║
    ╚═══════════════════════════════════════════════════════════════╝
    """
    print(banner)


def check_python_version():
    """Check if Python version is 3.9 or higher"""
    print("Checking Python version...")
    version = sys.version_info
    if version.major < 3 or (version.major == 3 and version.minor < 9):
        print(f"❌ Error: Python 3.9+ required. You have {sys.version}")
        sys.exit(1)
    print(f"✅ Python {sys.version.split()[0]} detected")


def check_tesseract():
    """Check if Tesseract OCR is installed"""
    print("\nChecking Tesseract OCR installation...")
    try:
        result = subprocess.run(['tesseract', '--version'], 
                              capture_output=True, text=True)
        if result.returncode == 0:
            print("✅ Tesseract OCR is installed")
            return True
    except FileNotFoundError:
        pass
    
    print("⚠️  Tesseract OCR not found")
    print("\nTo install Tesseract:")
    
    system = platform.system()
    if system == "Darwin":  # macOS
        print("  macOS: brew install tesseract")
    elif system == "Linux":
        print("  Ubuntu/Debian: sudo apt-get install tesseract-ocr")
        print("  Fedora: sudo dnf install tesseract")
    elif system == "Windows":
        print("  Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki")
    
    return False


def create_virtual_environment():
    """Create virtual environment if it doesn't exist"""
    print("\nSetting up virtual environment...")
    
    venv_path = Path("venv")
    if venv_path.exists():
        print("✅ Virtual environment already exists")
        return
    
    try:
        subprocess.run([sys.executable, "-m", "venv", "venv"], check=True)
        print("✅ Virtual environment created")
    except subprocess.CalledProcessError as e:
        print(f"❌ Failed to create virtual environment: {e}")
        sys.exit(1)


def get_pip_command():
    """Get the appropriate pip command for the virtual environment"""
    system = platform.system()
    if system == "Windows":
        pip_path = Path("venv/Scripts/pip.exe")
    else:
        pip_path = Path("venv/bin/pip")
    
    if pip_path.exists():
        return str(pip_path)
    else:
        return "pip"


def install_dependencies():
    """Install Python dependencies"""
    print("\nInstalling Python dependencies...")
    
    pip_cmd = get_pip_command()
    
    # Upgrade pip first
    print("Upgrading pip...")
    subprocess.run([pip_cmd, "install", "--upgrade", "pip"], 
                  capture_output=True)
    
    # Install requirements
    print("Installing requirements (this may take a few minutes)...")
    try:
        result = subprocess.run([pip_cmd, "install", "-r", "requirements.txt"],
                              capture_output=True, text=True)
        if result.returncode == 0:
            print("✅ All dependencies installed successfully")
        else:
            print(f"⚠️  Some dependencies may have failed: {result.stderr}")
    except subprocess.CalledProcessError as e:
        print(f"❌ Failed to install dependencies: {e}")
        sys.exit(1)


def create_directories():
    """Create necessary directory structure"""
    print("\nCreating directory structure...")
    
    directories = [
        "data",
        "data/Sourcepull",
        "data/Round1",
        "data/Round2",
        "cache",
        "logs",
        "reports",
        "config",
        "docs",
        "tests"
    ]
    
    for dir_path in directories:
        Path(dir_path).mkdir(parents=True, exist_ok=True)
    
    print("✅ Directory structure created")


def create_default_config():
    """Create default configuration file"""
    print("\nCreating default configuration...")
    
    config_path = Path("config/config.yaml")
    if config_path.exists():
        print("✅ Configuration file already exists")
        return
    
    config_content = """# Stanford Law Review Editorial Automation System Configuration

paths:
  master_sheet: "./data/Master_Sheet.xlsx"
  article_docx: "./data/article.docx"
  sourcepull_folder: "./data/Sourcepull/"
  round1_folder: "./data/Round1/"
  round2_folder: "./data/Round2/"
  cache_dir: "./cache/"
  output_dir: "./data/Sourcepull/"

apis:
  # Get your API keys from:
  # CourtListener: https://www.courtlistener.com/register/
  # GovInfo: https://api.govinfo.gov/signup
  courtlistener: ""
  crossref: ""  # Use your email
  govinfo: ""
  google_books: ""

preferences:
  remove_heinonline_covers: true
  remove_westlaw_headers: true
  max_retries: 3
  ocr_threshold: 0.7
  parallel_downloads: 5
  clean_pdf: true
  enable_ocr: true
  add_annotations: true
  min_quote_confidence: 0.85
  fuzzy_threshold: 80

bluebook:
  check_abbreviations: true
  check_punctuation: true
  check_signals: true
  check_short_forms: true
  check_capitalization: true
"""
    
    config_path.write_text(config_content)
    print("✅ Default configuration created at config/config.yaml")
    print("   ⚠️  Remember to add your API keys to the config file!")


def create_sample_files():
    """Create sample Excel and Word files for testing"""
    print("\nCreating sample files...")
    
    # Create sample Excel file
    try:
        import pandas as pd
        
        # Sourcepull sheet
        sourcepull_data = {
            'Source ID': ['001', '002', '003'],
            'Short Name': ['Smith_v_Jones', 'USC_Title_42', 'Harvard_Article'],
            'Full Citation': [
                'Smith v. Jones, 123 U.S. 456 (2020)',
                '42 U.S.C. § 1983 (2018)',
                'John Doe, Article Title, 100 Harv. L. Rev. 123 (2020)'
            ],
            'Completed?': ['', '', ''],
            'Location': ['', '', ''],
            'File Name': ['', '', ''],
            'Problems/Comments': ['', '', '']
        }
        
        # CC Round 1 sheet
        cc_data = {
            'Footnote #': [1, 1, 2],
            'Cite Order': [1, 2, 1],
            'Source ID': ['001', '002', '003'],
            'Quote': ['"sample quote"', '', '"another quote"'],
            'Supported?': ['', '', ''],
            'Issues': ['', '', ''],
            'Notes': ['', '', '']
        }
        
        # Create Excel file with multiple sheets
        excel_path = Path("data/Sample_Master_Sheet.xlsx")
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            pd.DataFrame(sourcepull_data).to_excel(writer, sheet_name='Sourcepull', index=False)
            pd.DataFrame(cc_data).to_excel(writer, sheet_name='CC_Round1', index=False)
        
        print(f"✅ Sample Excel file created at {excel_path}")
    except ImportError:
        print("⚠️  Could not create sample Excel file (pandas/openpyxl not installed)")
    
    # Create sample article
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Sample Law Review Article', 0)
        
        doc.add_paragraph('This is a sample paragraph with a citation.')
        doc.add_paragraph('', style='Normal')  # Space for footnote
        
        # Add sample footnotes (simplified)
        doc.add_paragraph('Footnotes:')
        doc.add_paragraph('1. Smith v. Jones, 123 U.S. 456, 458 (2020) ("This is a sample quote from the case.").')
        doc.add_paragraph('2. 42 U.S.C. § 1983 (2018).')
        doc.add_paragraph('3. John Doe, Article Title, 100 Harv. L. Rev. 123, 125 (2020).')
        
        doc_path = Path("data/Sample_Article.docx")
        doc.save(str(doc_path))
        
        print(f"✅ Sample Word document created at {doc_path}")
    except ImportError:
        print("⚠️  Could not create sample Word file (python-docx not installed)")


def create_scripts():
    """Create helper scripts"""
    print("\nCreating helper scripts...")
    
    # Create run script for Unix-like systems
    if platform.system() != "Windows":
        run_script = """#!/bin/bash
# Run SLRinator

# Activate virtual environment
source venv/bin/activate

# Run main program
python main.py "$@"
"""
        run_path = Path("run.sh")
        run_path.write_text(run_script)
        run_path.chmod(0o755)
        print("✅ Created run.sh script")
    
    # Create batch file for Windows
    else:
        run_script = """@echo off
REM Run SLRinator

REM Activate virtual environment
call venv\\Scripts\\activate

REM Run main program
python main.py %*
"""
        run_path = Path("run.bat")
        run_path.write_text(run_script)
        print("✅ Created run.bat script")


def print_next_steps():
    """Print next steps for the user"""
    print("\n" + "="*60)
    print("Setup Complete! 🎉")
    print("="*60)
    
    print("\n📝 Next Steps:")
    print("\n1. Add your API keys to config/config.yaml:")
    print("   - CourtListener: https://www.courtlistener.com/register/")
    print("   - GovInfo: https://api.govinfo.gov/signup")
    
    print("\n2. Place your files in the data/ directory:")
    print("   - Master Sheet: data/Master_Sheet.xlsx")
    print("   - Article: data/article.docx")
    
    print("\n3. Run the system:")
    if platform.system() == "Windows":
        print("   .\\run.bat --help           # Show help")
        print("   .\\run.bat setup            # Initialize system")
        print("   .\\run.bat sourcepull       # Run sourcepull")
        print("   .\\run.bat run --stage all  # Run full pipeline")
    else:
        print("   ./run.sh --help           # Show help")
        print("   ./run.sh setup            # Initialize system")
        print("   ./run.sh sourcepull       # Run sourcepull")
        print("   ./run.sh run --stage all  # Run full pipeline")
    
    print("\n4. Check the reports/ directory for results")
    
    print("\n📚 Documentation:")
    print("   - README.md for detailed instructions")
    print("   - config/config.yaml for configuration options")
    
    if not check_tesseract():
        print("\n⚠️  Remember to install Tesseract OCR for PDF text extraction")
    
    print("\n💡 Tip: Start with the sample files to test the system:")
    print("   - data/Sample_Master_Sheet.xlsx")
    print("   - data/Sample_Article.docx")


def main():
    """Main setup function"""
    print_banner()
    
    # Run setup steps
    check_python_version()
    has_tesseract = check_tesseract()
    create_virtual_environment()
    install_dependencies()
    create_directories()
    create_default_config()
    create_sample_files()
    create_scripts()
    
    # Print completion message
    print_next_steps()
    
    print("\n✨ Happy editing! The Stanford Law Review Editorial Team")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⚠️  Setup interrupted by user")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ Setup failed with error: {e}")
        sys.exit(1)