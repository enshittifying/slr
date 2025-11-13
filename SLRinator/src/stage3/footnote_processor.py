"""
Footnote Processor Module for Stanford Law Review
Extracts and processes footnotes from Word documents
"""

import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
import docx
from docx.document import Document
from docx.oxml.ns import qn
from docx.oxml import CT_P, CT_R

logger = logging.getLogger(__name__)


@dataclass
class Footnote:
    """Represents a footnote from the article"""
    number: int
    text: str
    citations: List[Dict[str, Any]] = field(default_factory=list)
    quotes: List[str] = field(default_factory=list)
    propositions: List[str] = field(default_factory=list)
    signals: List[str] = field(default_factory=list)
    parentheticals: List[str] = field(default_factory=list)
    reference_text: str = ""  # Text in main document that footnote supports
    page_num: Optional[int] = None
    
    def __post_init__(self):
        """Extract components from footnote text"""
        if self.text:
            self.extract_components()
    
    def extract_components(self):
        """Extract quotes, signals, and parentheticals from footnote text"""
        # Extract quotes (text within quotation marks)
        quote_pattern = r'"([^"]+)"'
        self.quotes = re.findall(quote_pattern, self.text)
        
        # Extract signals
        signal_patterns = [
            r'\b(See generally|See also|See|Cf\.|Compare|Contra|But see|But cf\.|See, e\.g\.,|E\.g\.,|Accord)\b'
        ]
        for pattern in signal_patterns:
            matches = re.findall(pattern, self.text, re.IGNORECASE)
            self.signals.extend(matches)
        
        # Extract parentheticals
        paren_pattern = r'\([^)]+\)'
        all_parens = re.findall(paren_pattern, self.text)
        
        # Filter out year parentheticals and court/date combinations
        for paren in all_parens:
            content = paren[1:-1]  # Remove parentheses
            # Check if it's not just a year or court/year
            if not re.match(r'^\d{4}$', content) and \
               not re.match(r'^[A-Z][^)]*\d{4}$', content):
                self.parentheticals.append(content)
    
    def get_citation_positions(self) -> List[Tuple[int, int, str]]:
        """Get positions of citations within footnote text
        Returns: List of (start, end, citation_text) tuples
        """
        positions = []
        
        # Common citation patterns
        patterns = [
            # Cases
            r'\b[A-Z][a-z]+\s+v\.\s+[A-Z][a-z]+[^,;.]*?(?:\d{4}|\))',
            # Statutes
            r'\d+\s+U\.S\.C\.[^;.]+',
            r'\d+\s+C\.F\.R\.[^;.]+',
            # Articles
            r'[A-Z][a-z]+,\s+[^,]+?,\s+\d+\s+[A-Z][^,]+?\s+\d+[^;.]*\(\d{4}\)',
            # Id. and supra references
            r'\bId\.\s*(?:at\s+\d+)?',
            r'\b[A-Z][a-z]+,?\s+supra\s+note\s+\d+',
        ]
        
        for pattern in patterns:
            for match in re.finditer(pattern, self.text):
                positions.append((match.start(), match.end(), match.group()))
        
        # Sort by position
        positions.sort(key=lambda x: x[0])
        
        return positions


class FootnoteProcessor:
    """Processes footnotes from Word documents"""
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.footnotes: List[Footnote] = []
        self.main_text_segments: List[Dict[str, Any]] = []
        
    def extract_footnotes(self, docx_path: str) -> List[Footnote]:
        """Extract all footnotes from a Word document"""
        docx_path = Path(docx_path)
        
        if not docx_path.exists():
            logger.error(f"Document not found: {docx_path}")
            return []
        
        try:
            doc = docx.Document(str(docx_path))
            
            # Extract footnotes using XML parsing
            footnotes = self._extract_footnotes_xml(doc)
            
            # Also extract main text segments for context
            self._extract_main_text_with_references(doc)
            
            # Match footnotes with their reference text
            self._match_footnotes_to_text()
            
            self.footnotes = footnotes
            return footnotes
            
        except Exception as e:
            logger.error(f"Error extracting footnotes: {e}")
            return []
    
    def _extract_footnotes_xml(self, doc: Document) -> List[Footnote]:
        """Extract footnotes using XML parsing for better accuracy"""
        footnotes = []
        
        # Access the document's XML
        doc_xml = doc._element
        
        # Find all footnote elements
        footnote_elements = doc_xml.xpath('//w:footnote', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        
        for fn_elem in footnote_elements:
            # Skip separator and continuation footnotes
            fn_type = fn_elem.get(qn('w:type'))
            if fn_type in ['separator', 'continuationSeparator']:
                continue
            
            # Get footnote ID
            fn_id = fn_elem.get(qn('w:id'))
            if not fn_id or fn_id in ['0', '-1']:  # Skip special footnotes
                continue
            
            # Extract footnote text
            text_elements = fn_elem.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            fn_text = ''.join(elem.text for elem in text_elements if elem.text)
            
            # Create Footnote object
            try:
                footnote = Footnote(
                    number=int(fn_id),
                    text=fn_text.strip()
                )
                footnotes.append(footnote)
            except ValueError:
                logger.warning(f"Could not parse footnote ID: {fn_id}")
        
        # Sort by number
        footnotes.sort(key=lambda x: x.number)
        
        # If XML extraction fails, fall back to paragraph-based extraction
        if not footnotes:
            footnotes = self._extract_footnotes_fallback(doc)
        
        return footnotes
    
    def _extract_footnotes_fallback(self, doc: Document) -> List[Footnote]:
        """Fallback method to extract footnotes from paragraphs"""
        footnotes = []
        footnote_pattern = r'^(\d+)\.\s+(.+)$'
        
        for para in doc.paragraphs:
            # Check if paragraph looks like a footnote
            match = re.match(footnote_pattern, para.text)
            if match:
                fn_num = int(match.group(1))
                fn_text = match.group(2)
                
                footnote = Footnote(
                    number=fn_num,
                    text=fn_text
                )
                footnotes.append(footnote)
        
        return footnotes
    
    def _extract_main_text_with_references(self, doc: Document):
        """Extract main text segments that contain footnote references"""
        self.main_text_segments = []
        
        for para_idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            
            # Look for footnote references (superscript numbers)
            # This is simplified - in practice would need XML parsing for accuracy
            fn_refs = self._find_footnote_references(para)
            
            if fn_refs:
                segment = {
                    'paragraph_index': para_idx,
                    'text': para.text,
                    'footnote_refs': fn_refs,
                    'page_estimate': para_idx // 3  # Rough page estimate
                }
                self.main_text_segments.append(segment)
    
    def _find_footnote_references(self, paragraph) -> List[int]:
        """Find footnote reference numbers in a paragraph"""
        refs = []
        
        # Access paragraph XML to find footnote references
        try:
            para_xml = paragraph._element
            fn_refs = para_xml.xpath('.//w:footnoteReference', 
                                    namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            for ref in fn_refs:
                fn_id = ref.get(qn('w:id'))
                if fn_id:
                    try:
                        refs.append(int(fn_id))
                    except ValueError:
                        pass
        except Exception as e:
            logger.debug(f"Could not extract footnote references: {e}")
        
        return refs
    
    def _match_footnotes_to_text(self):
        """Match footnotes to their reference text in the main document"""
        for segment in self.main_text_segments:
            for fn_num in segment['footnote_refs']:
                # Find corresponding footnote
                for footnote in self.footnotes:
                    if footnote.number == fn_num:
                        # Extract surrounding text as context
                        text = segment['text']
                        # Find sentence containing the footnote reference
                        sentences = re.split(r'[.!?]+', text)
                        
                        for sentence in sentences:
                            # Simplified check - would need more sophisticated matching
                            if str(fn_num) in sentence or f'[{fn_num}]' in sentence:
                                footnote.reference_text = sentence.strip()
                                footnote.page_num = segment['page_estimate']
                                break
    
    def parse_citations_in_footnotes(self) -> Dict[int, List[Dict[str, Any]]]:
        """Parse citations within each footnote"""
        results = {}
        
        for footnote in self.footnotes:
            citations = []
            positions = footnote.get_citation_positions()
            
            for idx, (start, end, cite_text) in enumerate(positions):
                citation = {
                    'order': idx + 1,
                    'text': cite_text,
                    'start_pos': start,
                    'end_pos': end,
                    'type': self._identify_citation_type(cite_text),
                    'is_id': 'Id.' in cite_text,
                    'is_supra': 'supra' in cite_text.lower(),
                    'is_infra': 'infra' in cite_text.lower()
                }
                
                # Extract pincites
                pincite_match = re.search(r'at\s+(\d+(?:[-–]\d+)?)', cite_text)
                if pincite_match:
                    citation['pincite'] = pincite_match.group(1)
                
                citations.append(citation)
            
            footnote.citations = citations
            results[footnote.number] = citations
        
        return results
    
    def _identify_citation_type(self, cite_text: str) -> str:
        """Identify the type of citation"""
        if ' v. ' in cite_text:
            return 'case'
        elif 'U.S.C.' in cite_text or 'C.F.R.' in cite_text:
            return 'statute'
        elif re.search(r'\d+\s+[A-Z].*\s+\d+', cite_text):
            return 'article'
        elif 'http' in cite_text or 'www.' in cite_text:
            return 'web'
        elif 'Id.' in cite_text:
            return 'id_reference'
        elif 'supra' in cite_text:
            return 'supra_reference'
        else:
            return 'unknown'
    
    def extract_quotes_from_footnotes(self) -> Dict[int, List[Dict[str, Any]]]:
        """Extract and analyze quotes from footnotes"""
        results = {}
        
        for footnote in self.footnotes:
            quotes = []
            
            for quote_text in footnote.quotes:
                quote_info = {
                    'text': quote_text,
                    'length': len(quote_text),
                    'has_ellipsis': '...' in quote_text or '. . .' in quote_text,
                    'has_brackets': '[' in quote_text and ']' in quote_text,
                    'has_emphasis': '(emphasis added)' in footnote.text.lower(),
                    'surrounding_text': self._get_quote_context(footnote.text, quote_text)
                }
                
                # Check for alterations
                alterations = []
                if quote_info['has_ellipsis']:
                    alterations.append('ellipsis')
                if quote_info['has_brackets']:
                    alterations.append('brackets')
                if quote_info['has_emphasis']:
                    alterations.append('emphasis')
                
                quote_info['alterations'] = alterations
                quotes.append(quote_info)
            
            results[footnote.number] = quotes
        
        return results
    
    def _get_quote_context(self, text: str, quote: str, context_chars: int = 50) -> str:
        """Get surrounding context for a quote"""
        quote_pos = text.find(quote)
        if quote_pos == -1:
            return ""
        
        start = max(0, quote_pos - context_chars)
        end = min(len(text), quote_pos + len(quote) + context_chars)
        
        context = text[start:end]
        
        # Add ellipsis if truncated
        if start > 0:
            context = "..." + context
        if end < len(text):
            context = context + "..."
        
        return context
    
    def identify_propositions(self) -> Dict[int, List[str]]:
        """Identify propositions that footnotes support"""
        results = {}
        
        for footnote in self.footnotes:
            propositions = []
            
            # The proposition is typically the sentence(s) before the footnote
            if footnote.reference_text:
                # Clean and process the reference text
                prop_text = footnote.reference_text
                
                # Remove the footnote number if present
                prop_text = re.sub(r'\[\d+\]|\d+$', '', prop_text).strip()
                
                propositions.append(prop_text)
            
            # Also check for explicit proposition statements in the footnote
            # (e.g., "arguing that...", "noting that...", "finding that...")
            proposition_patterns = [
                r'(?:arguing|noting|finding|stating|holding|concluding)\s+that\s+([^.]+)',
                r'for the proposition that\s+([^.]+)',
            ]
            
            for pattern in proposition_patterns:
                matches = re.findall(pattern, footnote.text, re.IGNORECASE)
                propositions.extend(matches)
            
            footnote.propositions = propositions
            results[footnote.number] = propositions
        
        return results
    
    def generate_citecheck_files(self, output_dir: str, round_num: int = 1) -> List[str]:
        """Generate files for citechecking round"""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        generated_files = []
        
        for footnote in self.footnotes:
            for cite_idx, citation in enumerate(footnote.citations):
                # Generate filename
                # Format: R[round]-[footnote:03d]-[cite_order:02d]-[source_id:03d]-[short_name].pdf
                filename = f"R{round_num}-{footnote.number:03d}-{cite_idx+1:02d}-XXX-temp.txt"
                filepath = output_dir / filename
                
                # Create temporary file with citation info
                content = f"""Footnote {footnote.number}
Citation {cite_idx + 1}: {citation['text']}
Type: {citation['type']}

Reference Text:
{footnote.reference_text}

Quotes to verify:
"""
                for quote in footnote.quotes:
                    content += f"- \"{quote}\"\n"
                
                content += f"\nPropositions to verify:\n"
                for prop in footnote.propositions:
                    content += f"- {prop}\n"
                
                filepath.write_text(content, encoding='utf-8')
                generated_files.append(str(filepath))
        
        return generated_files
    
    def export_footnotes_to_excel(self, output_path: str):
        """Export footnotes to Excel for review"""
        import pandas as pd
        
        data = []
        for footnote in self.footnotes:
            for cite_idx, citation in enumerate(footnote.citations):
                row = {
                    'Footnote #': footnote.number,
                    'Citation Order': cite_idx + 1,
                    'Citation Text': citation['text'],
                    'Citation Type': citation['type'],
                    'Quotes': '; '.join(footnote.quotes),
                    'Propositions': '; '.join(footnote.propositions),
                    'Reference Text': footnote.reference_text,
                    'Page Estimate': footnote.page_num
                }
                data.append(row)
        
        df = pd.DataFrame(data)
        df.to_excel(output_path, index=False)
        logger.info(f"Exported {len(data)} footnote citations to {output_path}")
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get statistics about processed footnotes"""
        total_quotes = sum(len(fn.quotes) for fn in self.footnotes)
        total_citations = sum(len(fn.citations) for fn in self.footnotes)
        
        citation_types = {}
        for fn in self.footnotes:
            for cite in fn.citations:
                cite_type = cite.get('type', 'unknown')
                citation_types[cite_type] = citation_types.get(cite_type, 0) + 1
        
        return {
            'total_footnotes': len(self.footnotes),
            'total_citations': total_citations,
            'total_quotes': total_quotes,
            'avg_citations_per_footnote': total_citations / len(self.footnotes) if self.footnotes else 0,
            'avg_quotes_per_footnote': total_quotes / len(self.footnotes) if self.footnotes else 0,
            'citation_types': citation_types,
            'footnotes_with_quotes': sum(1 for fn in self.footnotes if fn.quotes),
            'footnotes_with_propositions': sum(1 for fn in self.footnotes if fn.propositions)
        }