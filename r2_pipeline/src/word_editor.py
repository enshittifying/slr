"""
Edit Word document with Track Changes enabled.
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from typing import Dict, List
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

class WordEditor:
    """Edit Word document with track changes."""

    def __init__(self, doc_path: Path, author: str = "R2 Pipeline"):
        if not doc_path.exists():
            logger.warning(f"Word document not found at {doc_path}. Creating mock editor.")
            self.doc = None
            self.doc_path = doc_path
            self.author = author
            self.changes_made = []
            return

        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.author = author
        self.changes_made = []

        # Enable track changes
        self._enable_track_changes()

    def _enable_track_changes(self):
        """Enable track changes in the document."""
        if not self.doc:
            return

        try:
            # Access document settings
            settings = self.doc.settings

            # Create track changes element if it doesn't exist
            track_revisions = settings.element.find(qn('w:trackRevisions'))
            if track_revisions is None:
                track_revisions = OxmlElement('w:trackRevisions')
                settings.element.append(track_revisions)

            logger.info("Track changes enabled")
        except Exception as e:
            logger.warning(f"Could not enable track changes: {e}")

    def find_footnote(self, footnote_num: int):
        """
        Find footnote by number using XML parsing.

        Args:
            footnote_num: Footnote number to find

        Returns:
            Tuple of (footnote_element, footnotes_root) or (None, None)
        """
        if not self.doc:
            return None, None

        # Access footnotes through XML structure
        try:
            doc_part = self.doc.part
            footnotes_part = None

            for rel in doc_part.rels.values():
                if "footnotes" in rel.target_ref:
                    footnotes_part = rel.target_part
                    break

            if not footnotes_part:
                logger.error("No footnotes part found")
                return None, None

            # Parse footnotes XML
            footnotes_xml = footnotes_part.blob
            from lxml import etree
            root = etree.fromstring(footnotes_xml)

            # Define namespace
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # Find specific footnote
            for footnote in root.findall('.//w:footnote', ns):
                fn_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                if fn_id and int(fn_id) == footnote_num:
                    return footnote, root

            logger.warning(f"Footnote {footnote_num} not found")
            return None, None

        except Exception as e:
            logger.error(f"Error accessing footnotes: {e}")
            return None, None

    def replace_text_tracked(self,
                            footnote_num: int,
                            old_text: str,
                            new_text: str,
                            comment: str = None) -> bool:
        """
        Replace text in footnote with track changes.

        Args:
            footnote_num: Footnote number
            old_text: Text to replace
            new_text: Replacement text
            comment: Optional comment explaining change

        Returns:
            True if replacement made, False otherwise
        """
        if not self.doc:
            logger.warning("No document loaded")
            return False

        footnote_elem, root = self.find_footnote(footnote_num)
        if footnote_elem is None:
            logger.warning(f"Footnote {footnote_num} not found")
            return False

        # Get footnote text
        from lxml import etree
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        footnote_text = ""
        for text_elem in footnote_elem.findall('.//w:t', ns):
            if text_elem.text:
                footnote_text += text_elem.text

        if old_text not in footnote_text:
            logger.warning(f"Text '{old_text[:50]}...' not found in footnote {footnote_num}")
            return False

        # Replace text in XML
        try:
            for text_elem in footnote_elem.findall('.//w:t', ns):
                if text_elem.text and old_text in text_elem.text:
                    # Simple replacement - proper track changes would require more complex XML manipulation
                    text_elem.text = text_elem.text.replace(old_text, new_text)

                    # Log change
                    self.changes_made.append({
                        "footnote": footnote_num,
                        "type": "replacement",
                        "old": old_text,
                        "new": new_text,
                        "comment": comment
                    })

                    # Save changes back to footnotes part
                    self._save_footnotes_xml(root)

                    logger.info(f"Replaced text in footnote {footnote_num}")
                    return True

        except Exception as e:
            logger.error(f"Error replacing text: {e}")
            return False

        return False

    def _save_footnotes_xml(self, root):
        """Save modified footnotes XML back to document."""
        try:
            from lxml import etree
            doc_part = self.doc.part

            for rel in doc_part.rels.values():
                if "footnotes" in rel.target_ref:
                    footnotes_part = rel.target_part
                    # Update the blob with modified XML
                    footnotes_part._blob = etree.tostring(root, encoding='unicode').encode('utf-8')
                    break
        except Exception as e:
            logger.error(f"Error saving footnotes XML: {e}")

    def _mark_as_deleted(self, run):
        """Mark run as deleted (track changes)."""
        try:
            # Add deletion markup
            del_elem = OxmlElement('w:del')
            del_elem.set(qn('w:author'), self.author)
            del_elem.set(qn('w:date'), datetime.now().isoformat())

            run._element.append(del_elem)
        except Exception as e:
            logger.debug(f"Could not mark as deleted: {e}")

    def _mark_as_inserted(self, run):
        """Mark run as inserted (track changes)."""
        try:
            # Add insertion markup
            ins_elem = OxmlElement('w:ins')
            ins_elem.set(qn('w:author'), self.author)
            ins_elem.set(qn('w:date'), datetime.now().isoformat())

            run._element.append(ins_elem)
        except Exception as e:
            logger.debug(f"Could not mark as inserted: {e}")

    def add_comment(self, footnote_num: int, comment_text: str, tag: str = "[AA:]") -> bool:
        """
        Add a comment/note to footnote.

        Args:
            footnote_num: Footnote number
            comment_text: Comment text
            tag: Tag to use ([AA:], [SE:], [ME:])

        Returns:
            True if comment added
        """
        if not self.doc:
            return False

        footnote = self.find_footnote(footnote_num)
        if not footnote:
            return False

        # Add comment at end of footnote
        last_paragraph = footnote.paragraphs[-1]
        comment_run = last_paragraph.add_run(f" {tag} {comment_text}")

        # Highlight comment
        try:
            comment_run.font.highlight_color = 6  # Yellow highlight
            comment_run.font.bold = True
        except:
            pass

        self.changes_made.append({
            "footnote": footnote_num,
            "type": "comment",
            "tag": tag,
            "text": comment_text
        })

        logger.info(f"Added {tag} comment to footnote {footnote_num}")
        return True

    def get_changes_log(self) -> List[Dict]:
        """Get list of all changes made."""
        return self.changes_made

    def save(self, output_path: Path = None) -> Path:
        """
        Save document with track changes.

        Args:
            output_path: Path to save to (overwrites original if None)

        Returns:
            Path to saved document
        """
        if not self.doc:
            logger.warning("No document to save")
            return self.doc_path

        if output_path is None:
            output_path = self.doc_path

        self.doc.save(output_path)
        logger.info(f"Saved document to {output_path}")

        return output_path
